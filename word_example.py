import docx

d = docx.Document('c:\\users\\al\documents\\demo.docx')
d.paragraphs

d.paragraphs[0].text

p = d.paragraphs[1].text

p.runs

p.runs[0].text
p.runs[1].text
p.runs[2].text
p.runs[3].text

p.runs[1].bold

p.runs[0].bold == None

p.runs[3].italic

p.runs[3].underline = True
p.runs[3].text = 'italic and underlined.'

d.save('c:\\demo2.docx')

p.style

p.style = 'Title'
d.save('c:\\demo3.docx')

d = docx.Document()
d.add_paragraph('Hello this is a paragraph.')

d.add_paragraph('This is another paragraph.')
d.save('c:\\demo4.docx')
p = d.paragraphs[0]
p.add_run('This is a new run.')

p.runs

p.runs[1].bold = True
d.save('c:\\demo5.docx')

docx.Document()
